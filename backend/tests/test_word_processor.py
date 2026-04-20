"""
Tests for Word Document Processor

Tests the WordProcessor class for text extraction, format preservation,
and file validation.
"""

import pytest
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.services.document_processor import (
    DocumentType,
    TextSegment,
)
from src.services.word_processor import WordProcessor


class TestWordProcessor:
    """Tests for WordProcessor."""
    
    def test_supported_extensions(self):
        """Test supported extensions."""
        processor = WordProcessor()
        assert processor.supported_extensions == ['.docx']
    
    def test_document_type(self):
        """Test document type."""
        processor = WordProcessor()
        assert processor.document_type == DocumentType.WORD
    
    def test_generate_output_filename(self):
        """Test output filename generation with datetime stamp."""
        import re
        processor = WordProcessor()

        filename = processor.generate_output_filename(Path("document.docx"))
        assert re.match(r"^document_\d{8}_\d{6}_vi\.docx$", filename)

        filename = processor.generate_output_filename(Path("report.docx"), "en")
        assert re.match(r"^report_\d{8}_\d{6}_en\.docx$", filename)


class TestWordProcessorExtraction:
    """Tests for WordProcessor text extraction."""
    
    @pytest.mark.asyncio
    async def test_extract_text_from_paragraphs(self):
        """Test extracting text from paragraphs."""
        processor = WordProcessor()
        
        # Create a test Word file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Hello World")
            doc.add_paragraph("This is a test")
            doc.save(f.name)
            file_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(file_path)
            
            assert len(segments) == 2
            texts = [s.text for s in segments]
            assert "Hello World" in texts
            assert "This is a test" in texts
            
            # Check metadata
            for segment in segments:
                assert segment.metadata.get("type") == "paragraph"
                assert "paragraph_idx" in segment.metadata
        finally:
            file_path.unlink()

    @pytest.mark.asyncio
    async def test_extract_text_from_tables(self):
        """Test extracting text from tables."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Cell A1"
            table.cell(0, 1).text = "Cell B1"
            table.cell(1, 0).text = "Cell A2"
            table.cell(1, 1).text = "Cell B2"
            doc.save(f.name)
            file_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(file_path)
            
            # Should have 4 table cell segments
            table_segments = [s for s in segments if s.metadata.get("type") == "table_cell"]
            assert len(table_segments) == 4
            
            texts = [s.text for s in table_segments]
            assert "Cell A1" in texts
            assert "Cell B1" in texts
            assert "Cell A2" in texts
            assert "Cell B2" in texts
        finally:
            file_path.unlink()
    
    @pytest.mark.asyncio
    async def test_extract_text_preserves_formatting_metadata(self):
        """Test that extraction preserves formatting metadata."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run("Bold Text")
            run.bold = True
            run.font.size = Pt(14)
            doc.save(f.name)
            file_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(file_path)
            
            assert len(segments) == 1
            segment = segments[0]
            
            # Check runs metadata
            runs = segment.metadata.get("runs", [])
            assert len(runs) == 1
            assert runs[0]["bold"] is True
            assert runs[0]["font_size"] == 14.0
        finally:
            file_path.unlink()


class TestWordProcessorValidation:
    """Tests for WordProcessor file validation."""
    
    @pytest.mark.asyncio
    async def test_validate_file_success(self):
        """Test validating a valid Word file."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Test content")
            doc.save(f.name)
            file_path = Path(f.name)
        
        try:
            is_valid, error = await processor.validate_file(file_path)
            assert is_valid is True
            assert error is None
        finally:
            file_path.unlink()
    
    @pytest.mark.asyncio
    async def test_validate_file_not_found(self):
        """Test validating a non-existent file."""
        processor = WordProcessor()
        
        is_valid, error = await processor.validate_file(Path("/nonexistent/file.docx"))
        
        assert is_valid is False
        assert "not found" in error.lower()
    
    @pytest.mark.asyncio
    async def test_validate_file_wrong_extension(self):
        """Test validating a file with wrong extension."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as f:
            f.write(b"test")
            file_path = Path(f.name)
        
        try:
            is_valid, error = await processor.validate_file(file_path)
            assert is_valid is False
            assert "unsupported" in error.lower()
        finally:
            file_path.unlink()
    
    @pytest.mark.asyncio
    async def test_validate_file_corrupted(self):
        """Test validating a corrupted file."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            f.write(b"not a valid docx file")
            file_path = Path(f.name)
        
        try:
            is_valid, error = await processor.validate_file(file_path)
            assert is_valid is False
            assert "corrupted" in error.lower() or "cannot be read" in error.lower()
        finally:
            file_path.unlink()


class TestWordProcessorWriteTranslated:
    """Tests for WordProcessor write_translated method."""
    
    @pytest.mark.asyncio
    async def test_write_translated_paragraphs(self):
        """Test writing translated text to paragraphs."""
        processor = WordProcessor()
        
        # Create a test Word file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Hello")
            doc.add_paragraph("World")
            doc.save(f.name)
            input_path = Path(f.name)
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)
        
        try:
            # Extract text
            segments = await processor.extract_text(input_path)
            
            # Create translations
            translations = ["Xin chào", "Thế giới"]
            
            # Write translated (with output_mode="replace" to replace original text)
            success = await processor.write_translated(
                input_path, segments, translations, output_path, output_mode="replace"
            )
            
            assert success is True
            assert output_path.exists()
            
            # Verify the translated content
            doc = Document(str(output_path))
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            assert "Xin chào" in texts
            assert "Thế giới" in texts
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()
    
    @pytest.mark.asyncio
    async def test_write_translated_preserves_bold(self):
        """Test that writing preserves bold formatting."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run("Bold Text")
            run.bold = True
            doc.save(f.name)
            input_path = Path(f.name)
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(input_path)
            translations = ["Văn bản đậm"]
            
            # Write translated (with output_mode="replace" to replace original text)
            success = await processor.write_translated(
                input_path, segments, translations, output_path, output_mode="replace"
            )
            
            assert success is True
            
            # Verify bold is preserved
            doc = Document(str(output_path))
            para = doc.paragraphs[0]
            assert para.runs[0].bold is True
            assert para.runs[0].text == "Văn bản đậm"
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()
    
    @pytest.mark.asyncio
    async def test_write_translated_table_cells(self):
        """Test writing translated text to table cells."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Hello"
            table.cell(0, 1).text = "World"
            doc.save(f.name)
            input_path = Path(f.name)
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(input_path)
            table_segments = [s for s in segments if s.metadata.get("type") == "table_cell"]
            
            translations = ["Xin chào", "Thế giới"]
            
            # Write translated (with output_mode="replace" to replace original text)
            success = await processor.write_translated(
                input_path, table_segments, translations, output_path, output_mode="replace"
            )
            
            assert success is True
            
            # Verify table content
            doc = Document(str(output_path))
            table = doc.tables[0]
            assert table.cell(0, 0).text == "Xin chào"
            assert table.cell(0, 1).text == "Thế giới"
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()


class TestWordProcessorTOCParagraphs:
    """Tests for WordProcessor handling of TOC (Table of Contents) paragraphs."""

    @pytest.mark.asyncio
    async def test_write_translated_replaces_toc_hyperlink_paragraphs(self):
        """Test that TOC paragraphs with hyperlinks are fully replaced, not appended."""
        from lxml import etree
        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            # Add a normal paragraph first so the TOC paragraph gets a known index
            doc.add_paragraph("Intro")

            # Add a paragraph that simulates a TOC entry with a w:hyperlink element
            toc_para = doc.add_paragraph()
            toc_para_elem = toc_para._element
            nsmap_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

            # Build a hyperlink element containing a run with text (like a real TOC entry)
            hyperlink = etree.SubElement(toc_para_elem, f'{{{nsmap_w}}}hyperlink')
            hyperlink.set(f'{{{nsmap_w}}}anchor', '_Toc123')
            h_run = etree.SubElement(hyperlink, f'{{{nsmap_w}}}r')
            h_text = etree.SubElement(h_run, f'{{{nsmap_w}}}t')
            h_text.text = '目的与范围'

            doc.save(f.name)
            input_path = Path(f.name)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)

        try:
            segments = await processor.extract_text(input_path)
            # Find the TOC segment
            toc_segments = [s for s in segments if '目的与范围' in s.text]
            assert len(toc_segments) == 1

            translations = []
            all_segments = []
            for seg in segments:
                all_segments.append(seg)
                if '目的与范围' in seg.text:
                    translations.append('Purpose and Scope')
                else:
                    translations.append(seg.text)

            success = await processor.write_translated(
                input_path, all_segments, translations, output_path, output_mode="replace"
            )
            assert success is True

            # Verify the TOC paragraph only contains the translated text, not both
            result_doc = Document(str(output_path))
            toc_para_text = result_doc.paragraphs[1].text
            assert toc_para_text == 'Purpose and Scope', (
                f"Expected only translated text but got: '{toc_para_text}'"
            )
            # Ensure original text is not present
            assert '目的与范围' not in toc_para_text
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()


class TestWordProcessorMergedCells:
    """Tests for WordProcessor handling of merged cells."""
    
    @pytest.mark.asyncio
    async def test_extract_merged_cells_horizontal(self):
        """Test that horizontally merged cells are extracted only once."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=2, cols=3)
            
            # Merge cells A1 and B1 (horizontal merge)
            a1 = table.cell(0, 0)
            b1 = table.cell(0, 1)
            a1.merge(b1)
            a1.text = "Merged A1-B1"
            
            table.cell(0, 2).text = "C1"
            table.cell(1, 0).text = "A2"
            table.cell(1, 1).text = "B2"
            table.cell(1, 2).text = "C2"
            doc.save(f.name)
            file_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(file_path)
            table_segments = [s for s in segments if s.metadata.get("type") == "table_cell"]
            
            # Should have 5 segments (merged cell counts as 1, plus 4 other cells)
            assert len(table_segments) == 5
            
            texts = [s.text for s in table_segments]
            assert "Merged A1-B1" in texts
            assert texts.count("Merged A1-B1") == 1  # Should appear only once
            assert "C1" in texts
            assert "A2" in texts
            assert "B2" in texts
            assert "C2" in texts
        finally:
            file_path.unlink()
    
    @pytest.mark.asyncio
    async def test_write_translated_merged_cells(self):
        """Test writing translations to merged cells."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=3)
            
            # Merge cells A1 and B1
            a1 = table.cell(0, 0)
            b1 = table.cell(0, 1)
            a1.merge(b1)
            a1.text = "Merged"
            table.cell(0, 2).text = "Single"
            doc.save(f.name)
            input_path = Path(f.name)
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(input_path)
            table_segments = [s for s in segments if s.metadata.get("type") == "table_cell"]
            
            # Create translations for each segment
            translations = []
            for seg in table_segments:
                if seg.text == "Merged":
                    translations.append("Đã hợp nhất")
                elif seg.text == "Single":
                    translations.append("Đơn lẻ")
            
            # Write translated (with output_mode="replace" to replace original text)
            success = await processor.write_translated(
                input_path, table_segments, translations, output_path, output_mode="replace"
            )
            
            assert success is True
            
            # Verify content
            doc = Document(str(output_path))
            table = doc.tables[0]
            # Merged cell should have translated text
            assert "Đã hợp nhất" in table.cell(0, 0).text
            assert table.cell(0, 2).text == "Đơn lẻ"
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()


class TestWordProcessorNestedTables:
    """Tests for WordProcessor handling of nested tables."""
    
    @pytest.mark.asyncio
    async def test_extract_nested_tables(self):
        """Test extracting text from nested tables."""
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            
            # Create outer table
            outer_table = doc.add_table(rows=1, cols=2)
            outer_table.cell(0, 0).text = "Outer Cell"
            
            # Add nested table in second cell
            inner_cell = outer_table.cell(0, 1)
            # Clear default paragraph and add nested table
            inner_cell.paragraphs[0].clear()
            nested_table = inner_cell.add_table(rows=1, cols=2)
            nested_table.cell(0, 0).text = "Nested A"
            nested_table.cell(0, 1).text = "Nested B"
            
            doc.save(f.name)
            file_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(file_path)
            table_segments = [s for s in segments if s.metadata.get("type") == "table_cell"]
            
            texts = [s.text for s in table_segments]
            
            # Should extract from both outer and nested tables
            assert "Outer Cell" in texts
            assert "Nested A" in texts
            assert "Nested B" in texts
            
            # Check that nested table segments have correct location prefix
            nested_segments = [s for s in table_segments if "nested" in s.location]
            assert len(nested_segments) == 2
        finally:
            file_path.unlink()


class TestWordProcessorHeaderFooterTables:
    """Tests for WordProcessor handling of tables in headers/footers."""
    
    @pytest.mark.asyncio
    async def test_extract_table_in_header(self):
        """Test extracting text from tables in document header."""
        from docx.shared import Inches
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            
            # Add content to body
            doc.add_paragraph("Body content")
            
            # Add table to header (requires width parameter)
            section = doc.sections[0]
            header = section.header
            header_table = header.add_table(rows=1, cols=2, width=Inches(6))
            header_table.cell(0, 0).text = "Header Cell 1"
            header_table.cell(0, 1).text = "Header Cell 2"
            
            doc.save(f.name)
            file_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(file_path)
            
            # Check for header table segments
            header_table_segments = [
                s for s in segments 
                if s.metadata.get("type") == "table_cell" 
                and s.metadata.get("header_footer") == "header"
            ]
            
            assert len(header_table_segments) == 2
            texts = [s.text for s in header_table_segments]
            assert "Header Cell 1" in texts
            assert "Header Cell 2" in texts
        finally:
            file_path.unlink()
    
    @pytest.mark.asyncio
    async def test_extract_table_in_footer(self):
        """Test extracting text from tables in document footer."""
        from docx.shared import Inches
        processor = WordProcessor()
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            
            # Add content to body
            doc.add_paragraph("Body content")
            
            # Add table to footer (requires width parameter)
            section = doc.sections[0]
            footer = section.footer
            footer_table = footer.add_table(rows=1, cols=2, width=Inches(6))
            footer_table.cell(0, 0).text = "Footer Cell 1"
            footer_table.cell(0, 1).text = "Footer Cell 2"
            
            doc.save(f.name)
            file_path = Path(f.name)
        
        try:
            segments = await processor.extract_text(file_path)
            
            # Check for footer table segments
            footer_table_segments = [
                s for s in segments 
                if s.metadata.get("type") == "table_cell" 
                and s.metadata.get("header_footer") == "footer"
            ]
            
            assert len(footer_table_segments) == 2
            texts = [s.text for s in footer_table_segments]
            assert "Footer Cell 1" in texts
            assert "Footer Cell 2" in texts
        finally:
            file_path.unlink()


class TestWordProcessorMultiParagraphCells:
    """Tests for multi-paragraph table cell write-back preservation."""

    @pytest.mark.asyncio
    async def test_write_translated_preserves_multi_paragraph_cells(self):
        """Test that multi-paragraph cell structure is preserved during write-back."""
        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            # Create two paragraphs in one cell
            cell.paragraphs[0].text = "Line 1"
            cell.add_paragraph("Line 2")
            doc.save(f.name)
            input_path = Path(f.name)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)

        try:
            segments = await processor.extract_text(input_path)
            table_segments = [s for s in segments if s.metadata.get("type") == "table_cell"]
            assert len(table_segments) == 1

            # Translation with newline to map to two paragraphs
            translations = ["Dòng 1\nDòng 2"]

            success = await processor.write_translated(
                input_path, table_segments, translations, output_path, output_mode="replace"
            )
            assert success is True

            result_doc = Document(str(output_path))
            cell = result_doc.tables[0].cell(0, 0)
            para_texts = [p.text for p in cell.paragraphs if p.text.strip()]
            assert para_texts == ["Dòng 1", "Dòng 2"]
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()


class TestWordProcessorVMLTextboxes:
    """Tests for VML legacy text box support."""

    @pytest.mark.asyncio
    async def test_extract_vml_textbox(self):
        """Test extracting text from VML text boxes."""
        from lxml import etree

        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Body text")

            # Inject a VML textbox into the document body XML
            body = doc.element.body
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            v_ns = 'urn:schemas-microsoft-com:vml'
            o_ns = 'urn:schemas-microsoft-com:office:office'

            # Create w:r > w:pict > v:shape > v:textbox > w:txbxContent > w:p > w:r > w:t
            p_elem = etree.SubElement(body, f'{{{w_ns}}}p')
            r_elem = etree.SubElement(p_elem, f'{{{w_ns}}}r')
            pict_elem = etree.SubElement(r_elem, f'{{{w_ns}}}pict')
            shape_elem = etree.SubElement(pict_elem, f'{{{v_ns}}}shape')
            textbox_elem = etree.SubElement(shape_elem, f'{{{v_ns}}}textbox')
            txbx_content = etree.SubElement(textbox_elem, f'{{{w_ns}}}txbxContent')
            inner_p = etree.SubElement(txbx_content, f'{{{w_ns}}}p')
            inner_r = etree.SubElement(inner_p, f'{{{w_ns}}}r')
            inner_t = etree.SubElement(inner_r, f'{{{w_ns}}}t')
            inner_t.text = 'VML Text'

            doc.save(f.name)
            file_path = Path(f.name)

        try:
            segments = await processor.extract_text(file_path)
            textbox_segments = [s for s in segments if s.metadata.get("type") == "textbox"]
            texts = [s.text for s in textbox_segments]
            assert "VML Text" in texts
        finally:
            file_path.unlink()


class TestWordProcessorTextboxTables:
    """Tests for tables inside text boxes."""

    @pytest.mark.asyncio
    async def test_extract_table_in_textbox(self):
        """Test extracting table content from inside a text box."""
        from lxml import etree

        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Body text")

            body = doc.element.body
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            wps_ns = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'

            # Create a DrawingML textbox with a table inside
            p_elem = etree.SubElement(body, f'{{{w_ns}}}p')
            r_elem = etree.SubElement(p_elem, f'{{{w_ns}}}r')

            # Build wps:txbx > w:txbxContent > w:tbl > w:tr > w:tc > w:p > w:r > w:t
            mc_ns = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
            wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'

            drawing = etree.SubElement(r_elem, f'{{{w_ns}}}drawing')
            inline = etree.SubElement(drawing, f'{{{wp_ns}}}inline')
            graphic = etree.SubElement(inline, f'{{{a_ns}}}graphic')
            graphic_data = etree.SubElement(graphic, f'{{{a_ns}}}graphicData')
            wsp = etree.SubElement(graphic_data, f'{{{wps_ns}}}wsp')
            txbx = etree.SubElement(wsp, f'{{{wps_ns}}}txbx')
            txbx_content = etree.SubElement(txbx, f'{{{w_ns}}}txbxContent')

            # Add a table inside the textbox
            tbl = etree.SubElement(txbx_content, f'{{{w_ns}}}tbl')
            tblPr = etree.SubElement(tbl, f'{{{w_ns}}}tblPr')
            tr = etree.SubElement(tbl, f'{{{w_ns}}}tr')
            tc = etree.SubElement(tr, f'{{{w_ns}}}tc')
            tc_p = etree.SubElement(tc, f'{{{w_ns}}}p')
            tc_r = etree.SubElement(tc_p, f'{{{w_ns}}}r')
            tc_t = etree.SubElement(tc_r, f'{{{w_ns}}}t')
            tc_t.text = 'Textbox Table Cell'

            doc.save(f.name)
            file_path = Path(f.name)

        try:
            segments = await processor.extract_text(file_path)
            tb_table_segments = [s for s in segments if s.metadata.get("type") == "textbox_table_cell"]
            texts = [s.text for s in tb_table_segments]
            assert "Textbox Table Cell" in texts
        finally:
            file_path.unlink()

    @pytest.mark.asyncio
    async def test_write_translated_textbox_table(self):
        """Test writing translation to a table inside a text box."""
        from lxml import etree

        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Body")

            body = doc.element.body
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            wps_ns = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
            wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'

            p_elem = etree.SubElement(body, f'{{{w_ns}}}p')
            r_elem = etree.SubElement(p_elem, f'{{{w_ns}}}r')
            drawing = etree.SubElement(r_elem, f'{{{w_ns}}}drawing')
            inline = etree.SubElement(drawing, f'{{{wp_ns}}}inline')
            graphic = etree.SubElement(inline, f'{{{a_ns}}}graphic')
            graphic_data = etree.SubElement(graphic, f'{{{a_ns}}}graphicData')
            wsp = etree.SubElement(graphic_data, f'{{{wps_ns}}}wsp')
            txbx = etree.SubElement(wsp, f'{{{wps_ns}}}txbx')
            txbx_content = etree.SubElement(txbx, f'{{{w_ns}}}txbxContent')

            tbl = etree.SubElement(txbx_content, f'{{{w_ns}}}tbl')
            etree.SubElement(tbl, f'{{{w_ns}}}tblPr')
            tr = etree.SubElement(tbl, f'{{{w_ns}}}tr')
            tc = etree.SubElement(tr, f'{{{w_ns}}}tc')
            tc_p = etree.SubElement(tc, f'{{{w_ns}}}p')
            tc_r = etree.SubElement(tc_p, f'{{{w_ns}}}r')
            tc_t = etree.SubElement(tc_r, f'{{{w_ns}}}t')
            tc_t.text = 'Original'

            doc.save(f.name)
            input_path = Path(f.name)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)

        try:
            segments = await processor.extract_text(input_path)
            tb_table_segments = [s for s in segments if s.metadata.get("type") == "textbox_table_cell"]
            assert len(tb_table_segments) == 1

            success = await processor.write_translated(
                input_path, tb_table_segments, ["Đã dịch"], output_path, output_mode="replace"
            )
            assert success is True

            # Verify by re-extracting
            processor2 = WordProcessor()
            segments2 = await processor2.extract_text(output_path)
            tb_segments2 = [s for s in segments2 if s.metadata.get("type") == "textbox_table_cell"]
            assert len(tb_segments2) == 1
            assert tb_segments2[0].text == "Đã dịch"
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()


class TestWordProcessorSDT:
    """Tests for SDT/content control support."""

    @pytest.mark.asyncio
    async def test_extract_sdt_paragraph(self):
        """Test extracting text from SDT content controls."""
        from lxml import etree

        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Normal paragraph")

            body = doc.element.body
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

            # Add an SDT element wrapping a paragraph
            sdt = etree.SubElement(body, f'{{{w_ns}}}sdt')
            sdt_pr = etree.SubElement(sdt, f'{{{w_ns}}}sdtPr')
            sdt_content = etree.SubElement(sdt, f'{{{w_ns}}}sdtContent')
            p = etree.SubElement(sdt_content, f'{{{w_ns}}}p')
            r = etree.SubElement(p, f'{{{w_ns}}}r')
            t = etree.SubElement(r, f'{{{w_ns}}}t')
            t.text = 'SDT Content'

            doc.save(f.name)
            file_path = Path(f.name)

        try:
            segments = await processor.extract_text(file_path)
            sdt_segments = [s for s in segments if s.metadata.get("type") == "sdt_paragraph"]
            assert len(sdt_segments) == 1
            assert sdt_segments[0].text == "SDT Content"
        finally:
            file_path.unlink()

    @pytest.mark.asyncio
    async def test_write_translated_sdt_paragraph(self):
        """Test writing translation to an SDT paragraph."""
        from lxml import etree

        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Normal")

            body = doc.element.body
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

            sdt = etree.SubElement(body, f'{{{w_ns}}}sdt')
            etree.SubElement(sdt, f'{{{w_ns}}}sdtPr')
            sdt_content = etree.SubElement(sdt, f'{{{w_ns}}}sdtContent')
            p = etree.SubElement(sdt_content, f'{{{w_ns}}}p')
            r = etree.SubElement(p, f'{{{w_ns}}}r')
            t = etree.SubElement(r, f'{{{w_ns}}}t')
            t.text = 'Original SDT'

            doc.save(f.name)
            input_path = Path(f.name)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)

        try:
            segments = await processor.extract_text(input_path)
            sdt_segments = [s for s in segments if s.metadata.get("type") == "sdt_paragraph"]
            assert len(sdt_segments) == 1

            success = await processor.write_translated(
                input_path, sdt_segments, ["SDT Đã dịch"], output_path, output_mode="replace"
            )
            assert success is True

            processor2 = WordProcessor()
            segments2 = await processor2.extract_text(output_path)
            sdt_segments2 = [s for s in segments2 if s.metadata.get("type") == "sdt_paragraph"]
            assert len(sdt_segments2) == 1
            assert sdt_segments2[0].text == "SDT Đã dịch"
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()

    @pytest.mark.asyncio
    async def test_extract_sdt_table(self):
        """Test extracting table content from inside an SDT."""
        from lxml import etree

        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Normal")

            body = doc.element.body
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

            sdt = etree.SubElement(body, f'{{{w_ns}}}sdt')
            etree.SubElement(sdt, f'{{{w_ns}}}sdtPr')
            sdt_content = etree.SubElement(sdt, f'{{{w_ns}}}sdtContent')

            tbl = etree.SubElement(sdt_content, f'{{{w_ns}}}tbl')
            etree.SubElement(tbl, f'{{{w_ns}}}tblPr')
            tr = etree.SubElement(tbl, f'{{{w_ns}}}tr')
            tc = etree.SubElement(tr, f'{{{w_ns}}}tc')
            p = etree.SubElement(tc, f'{{{w_ns}}}p')
            r = etree.SubElement(p, f'{{{w_ns}}}r')
            t = etree.SubElement(r, f'{{{w_ns}}}t')
            t.text = 'SDT Table Cell'

            doc.save(f.name)
            file_path = Path(f.name)

        try:
            segments = await processor.extract_text(file_path)
            sdt_table_segments = [s for s in segments if s.metadata.get("type") == "sdt_table_cell"]
            assert len(sdt_table_segments) == 1
            assert sdt_table_segments[0].text == "SDT Table Cell"
        finally:
            file_path.unlink()


def _create_doc_with_footnote(doc, footnote_text):
    """Helper to add a footnote part to a Document."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    footnotes_xml = (
        f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:footnotes xmlns:w="{w_ns}">'
        f'  <w:footnote w:type="separator" w:id="-1">'
        f'    <w:p><w:r><w:separator/></w:r></w:p>'
        f'  </w:footnote>'
        f'  <w:footnote w:type="continuationSeparator" w:id="0">'
        f'    <w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        f'  </w:footnote>'
        f'  <w:footnote w:id="1">'
        f'    <w:p><w:r><w:t>{footnote_text}</w:t></w:r></w:p>'
        f'  </w:footnote>'
        f'</w:footnotes>'
    ).encode('utf-8')

    footnote_part = Part(
        PackURI('/word/footnotes.xml'),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
        footnotes_xml,
        doc.part.package
    )
    doc.part.relate_to(
        footnote_part,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
    )


class TestWordProcessorFootnotes:
    """Tests for footnote and endnote support."""

    @pytest.mark.asyncio
    async def test_extract_footnotes(self):
        """Test extracting text from footnotes."""
        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Body with footnote reference")
            _create_doc_with_footnote(doc, "This is a footnote")
            doc.save(f.name)
            file_path = Path(f.name)

        try:
            segments = await processor.extract_text(file_path)
            fn_segments = [s for s in segments if s.metadata.get("type") == "footnote"]
            assert len(fn_segments) == 1
            assert fn_segments[0].text == "This is a footnote"
        finally:
            file_path.unlink()

    @pytest.mark.asyncio
    async def test_write_translated_footnote(self):
        """Test writing translation to a footnote."""
        processor = WordProcessor()

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc = Document()
            doc.add_paragraph("Body text")
            _create_doc_with_footnote(doc, "Original footnote")
            doc.save(f.name)
            input_path = Path(f.name)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = Path(f.name)

        try:
            segments = await processor.extract_text(input_path)
            fn_segments = [s for s in segments if s.metadata.get("type") == "footnote"]
            assert len(fn_segments) == 1

            success = await processor.write_translated(
                input_path, fn_segments, ["Chú thích đã dịch"], output_path, output_mode="replace"
            )
            assert success is True

            # Verify by re-extracting
            processor2 = WordProcessor()
            segments2 = await processor2.extract_text(output_path)
            fn_segments2 = [s for s in segments2 if s.metadata.get("type") == "footnote"]
            assert len(fn_segments2) == 1
            assert fn_segments2[0].text == "Chú thích đã dịch"
        finally:
            input_path.unlink()
            if output_path.exists():
                output_path.unlink()
