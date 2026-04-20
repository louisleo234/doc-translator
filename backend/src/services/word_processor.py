"""
Word Document Processor Module

Provides document processing for Microsoft Word (.docx) files.
Extracts text from paragraphs, tables, headers, footers, and text boxes
while preserving formatting during translation writing.

Handles:
- Merged cells in tables (avoids duplicate extraction)
- Nested tables (tables within table cells)
- Tables in headers/footers
- Tables in text boxes
"""

import logging
from pathlib import Path
from typing import List, Optional, Any, Set, Tuple
import asyncio

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml.ns import qn, nsmap
from docx.oxml import parse_xml

from .document_processor import (
    DocumentProcessor,
    DocumentType,
    TextSegment,
    apply_output_mode,
)


class WordProcessor(DocumentProcessor):
    """
    Document processor for Microsoft Word files (.docx).
    
    Extracts text from:
    - Body paragraphs
    - Tables (cell by cell)
    - Headers and footers
    - Text boxes and shapes
    
    Preserves formatting during translation writing:
    - Font properties (name, size, color, bold, italic, underline)
    - Paragraph formatting (alignment, spacing, indentation)
    - Table structure and cell formatting
    - Embedded images and shapes
    """
    
    def __init__(self, logger: Optional[logging.Logger] = None):
        """
        Initialize the Word document processor.
        
        Args:
            logger: Logger instance for logging operations
        """
        self.logger = logger or logging.getLogger(__name__)

    @property
    def supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        return ['.docx']
    
    @property
    def document_type(self) -> DocumentType:
        """Return the document type this processor handles."""
        return DocumentType.WORD
    
    async def extract_text(self, file_path: Path) -> List[TextSegment]:
        """
        Extract all translatable text segments from the Word document.
        
        Extracts text from:
        - Body paragraphs
        - Tables
        - Headers and footers
        - Text boxes
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            List of TextSegment objects containing text and metadata
            
        Raises:
            ValueError: If the file cannot be loaded
        """
        try:
            document = await asyncio.to_thread(Document, str(file_path))
        except PackageNotFoundError:
            raise ValueError(f"Failed to load Word file: {file_path}")
        except Exception as e:
            raise ValueError(f"Failed to load Word file: {file_path}. Error: {str(e)}")
        
        segments: List[TextSegment] = []
        segment_id = 0
        
        # Extract from body paragraphs
        for para_idx, paragraph in enumerate(document.paragraphs):
            if paragraph.text.strip():
                segment = self._create_paragraph_segment(
                    paragraph, para_idx, segment_id, "body"
                )
                segments.append(segment)
                segment_id += 1
        
        # Extract from tables
        for table_idx, table in enumerate(document.tables):
            table_segments = self._extract_table_segments(
                table, table_idx, segment_id
            )
            segments.extend(table_segments)
            segment_id += len(table_segments)
        
        # Extract from headers and footers
        header_footer_segments = self._extract_header_footer_segments(
            document, segment_id
        )
        segments.extend(header_footer_segments)
        segment_id += len(header_footer_segments)
        
        # Extract from text boxes
        textbox_segments = await self._extract_textbox_segments(
            document, segment_id
        )
        segments.extend(textbox_segments)
        segment_id += len(textbox_segments)

        # Extract from structured document tags (content controls)
        sdt_segments = self._extract_sdt_segments(document, segment_id)
        segments.extend(sdt_segments)
        segment_id += len(sdt_segments)

        # Extract from footnotes and endnotes
        footnote_segments = self._extract_footnote_endnote_segments(
            document, segment_id
        )
        segments.extend(footnote_segments)

        self.logger.info(f"Extracted {len(segments)} text segments from {file_path.name}")
        return segments
    
    def _create_paragraph_segment(
        self,
        paragraph: Paragraph,
        para_idx: int,
        segment_id: int,
        location_prefix: str
    ) -> TextSegment:
        """Create a TextSegment from a paragraph with run-level metadata."""
        runs_metadata = []
        for run_idx, run in enumerate(paragraph.runs):
            run_meta = {
                "run_idx": run_idx,
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None,
                "font_color": self._get_font_color(run),
            }
            runs_metadata.append(run_meta)
        
        return TextSegment(
            id=str(segment_id),
            text=paragraph.text,
            location=f"{location_prefix}_paragraph_{para_idx}",
            metadata={
                "type": "paragraph",
                "location_prefix": location_prefix,
                "paragraph_idx": para_idx,
                "alignment": str(paragraph.alignment) if paragraph.alignment else None,
                "runs": runs_metadata,
            }
        )
    
    def _get_font_color(self, run) -> Optional[str]:
        """Extract font color as hex string if available."""
        try:
            if run.font.color and run.font.color.rgb:
                return str(run.font.color.rgb)
        except Exception:
            pass
        return None

    def _extract_table_segments(
        self,
        table: Table,
        table_idx: int,
        start_segment_id: int,
        location_prefix: str = "",
        processed_cells: Optional[Set[int]] = None
    ) -> List[TextSegment]:
        """
        Extract text segments from a table, handling merged cells and nested tables.
        
        Args:
            table: The table to extract from
            table_idx: Index of the table
            start_segment_id: Starting segment ID
            location_prefix: Prefix for location (e.g., "header_0_" for tables in headers)
            processed_cells: Set of already processed cell IDs to avoid duplicates
            
        Returns:
            List of TextSegment objects
        """
        segments = []
        segment_id = start_segment_id
        
        # Track processed cells to handle merged cells
        # For merged cells, python-docx returns the same _tc element for multiple positions
        # We use (cell_id, cell_text) tuple to detect true duplicates (merged cells have same id AND text)
        # This avoids false positives from python-docx's internal object caching
        table_processed_cells: Set[Tuple[int, str]] = set()
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                # Use cell's XML element id + text to detect merged cells
                # Merged cells have the same _tc element AND same text content
                cell_id = id(cell._tc)
                cell_key = (cell_id, cell_text)
                
                # Skip if this exact cell (same id AND text) was already processed
                if cell_key in table_processed_cells:
                    continue
                table_processed_cells.add(cell_key)
                
                cell_text = cell.text.strip()
                if cell_text:
                    # Get paragraph-level metadata for the cell
                    cell_paragraphs_meta = []
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            runs_meta = []
                            for run_idx, run in enumerate(para.runs):
                                runs_meta.append({
                                    "run_idx": run_idx,
                                    "text": run.text,
                                    "bold": run.bold,
                                    "italic": run.italic,
                                    "underline": run.underline,
                                    "font_name": run.font.name,
                                    "font_size": run.font.size.pt if run.font.size else None,
                                    "font_color": self._get_font_color(run),
                                })
                            cell_paragraphs_meta.append({
                                "para_idx": para_idx,
                                "text": para.text,
                                "runs": runs_meta,
                            })
                    
                    table_location = f"{location_prefix}table_{table_idx}" if location_prefix else f"table_{table_idx}"
                    segment = TextSegment(
                        id=str(segment_id),
                        text=cell_text,
                        location=f"{table_location}_row_{row_idx}_col_{col_idx}",
                        metadata={
                            "type": "table_cell",
                            "table_idx": table_idx,
                            "row_idx": row_idx,
                            "col_idx": col_idx,
                            "location_prefix": location_prefix,
                            "paragraphs": cell_paragraphs_meta,
                        }
                    )
                    segments.append(segment)
                    segment_id += 1
                
                # Extract from nested tables within this cell
                nested_tables = cell.tables
                for nested_idx, nested_table in enumerate(nested_tables):
                    nested_prefix = f"{location_prefix}table_{table_idx}_row_{row_idx}_col_{col_idx}_nested_"
                    nested_segments = self._extract_table_segments(
                        nested_table,
                        nested_idx,
                        segment_id,
                        location_prefix=nested_prefix,
                        processed_cells=processed_cells
                    )
                    segments.extend(nested_segments)
                    segment_id += len(nested_segments)
        
        return segments
    
    def _extract_header_footer_segments(
        self,
        document: Document,
        start_segment_id: int
    ) -> List[TextSegment]:
        """Extract text segments from headers and footers, including tables."""
        segments = []
        segment_id = start_segment_id
        
        for section_idx, section in enumerate(document.sections):
            # Extract from header
            header = section.header
            if header and not header.is_linked_to_previous:
                # Extract paragraphs
                for para_idx, paragraph in enumerate(header.paragraphs):
                    if paragraph.text.strip():
                        segment = self._create_paragraph_segment(
                            paragraph, para_idx, segment_id,
                            f"section_{section_idx}_header"
                        )
                        segment.metadata["section_idx"] = section_idx
                        segment.metadata["header_footer"] = "header"
                        segments.append(segment)
                        segment_id += 1
                
                # Extract tables in header
                for table_idx, table in enumerate(header.tables):
                    table_segments = self._extract_table_segments(
                        table, table_idx, segment_id,
                        location_prefix=f"section_{section_idx}_header_"
                    )
                    for seg in table_segments:
                        seg.metadata["section_idx"] = section_idx
                        seg.metadata["header_footer"] = "header"
                    segments.extend(table_segments)
                    segment_id += len(table_segments)
            
            # Extract from footer
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                # Extract paragraphs
                for para_idx, paragraph in enumerate(footer.paragraphs):
                    if paragraph.text.strip():
                        segment = self._create_paragraph_segment(
                            paragraph, para_idx, segment_id,
                            f"section_{section_idx}_footer"
                        )
                        segment.metadata["section_idx"] = section_idx
                        segment.metadata["header_footer"] = "footer"
                        segments.append(segment)
                        segment_id += 1
                
                # Extract tables in footer
                for table_idx, table in enumerate(footer.tables):
                    table_segments = self._extract_table_segments(
                        table, table_idx, segment_id,
                        location_prefix=f"section_{section_idx}_footer_"
                    )
                    for seg in table_segments:
                        seg.metadata["section_idx"] = section_idx
                        seg.metadata["header_footer"] = "footer"
                    segments.extend(table_segments)
                    segment_id += len(table_segments)

        return segments

    def _extract_sdt_segments(
        self,
        document: Document,
        start_segment_id: int
    ) -> List[TextSegment]:
        """
        Extract text segments from Structured Document Tags (content controls).

        SDT elements (w:sdt) wrap content that document.paragraphs/tables skip.
        Only processes SDTs that are direct children of the body.
        """
        segments = []
        segment_id = start_segment_id
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        try:
            body = document.element.body

            # Collect IDs of elements already reachable via document.paragraphs/tables
            known_para_ids = {id(p._element) for p in document.paragraphs}
            known_tbl_ids = {id(t._tbl) for t in document.tables}

            # Find SDT elements that are direct children of body
            sdt_idx = 0
            for child in body:
                if child.tag != f'{{{w_ns}}}sdt':
                    continue

                sdt_content = child.find(f'{{{w_ns}}}sdtContent')
                if sdt_content is None:
                    sdt_idx += 1
                    continue

                # Extract paragraphs inside this SDT
                para_counter = 0
                for elem in sdt_content:
                    if elem.tag == f'{{{w_ns}}}p':
                        if id(elem) in known_para_ids:
                            continue
                        text_parts = []
                        for t in elem.findall(f'.//{{{w_ns}}}t'):
                            if t.text:
                                text_parts.append(t.text)
                        full_text = ''.join(text_parts).strip()
                        if full_text:
                            segment = TextSegment(
                                id=str(segment_id),
                                text=full_text,
                                location=f"sdt_{sdt_idx}_paragraph_{para_counter}",
                                metadata={
                                    "type": "sdt_paragraph",
                                    "sdt_idx": sdt_idx,
                                    "paragraph_idx": para_counter,
                                }
                            )
                            segments.append(segment)
                            segment_id += 1
                        para_counter += 1

                    elif elem.tag == f'{{{w_ns}}}tbl':
                        if id(elem) in known_tbl_ids:
                            continue
                        # Extract table cells at XML level
                        tbl_idx_in_sdt = 0
                        rows = elem.findall(f'{{{w_ns}}}tr')
                        for row_idx, row_elem in enumerate(rows):
                            cells = row_elem.findall(f'{{{w_ns}}}tc')
                            for col_idx, cell_elem in enumerate(cells):
                                cell_text_parts = []
                                for t in cell_elem.findall(f'.//{{{w_ns}}}t'):
                                    if t.text:
                                        cell_text_parts.append(t.text)
                                cell_text = ''.join(cell_text_parts).strip()
                                if cell_text:
                                    segment = TextSegment(
                                        id=str(segment_id),
                                        text=cell_text,
                                        location=f"sdt_{sdt_idx}_table_{tbl_idx_in_sdt}_row_{row_idx}_col_{col_idx}",
                                        metadata={
                                            "type": "sdt_table_cell",
                                            "sdt_idx": sdt_idx,
                                            "table_idx": tbl_idx_in_sdt,
                                            "row_idx": row_idx,
                                            "col_idx": col_idx,
                                        }
                                    )
                                    segments.append(segment)
                                    segment_id += 1
                        tbl_idx_in_sdt += 1

                sdt_idx += 1

        except Exception as e:
            self.logger.warning(f"Error extracting SDT content: {e}")

        return segments

    def _extract_footnote_endnote_segments(
        self,
        document: Document,
        start_segment_id: int
    ) -> List[TextSegment]:
        """Extract text segments from footnotes and endnotes."""
        segments = []
        segment_id = start_segment_id
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        for note_type in ("footnote", "endnote"):
            try:
                part_name = f"/word/{note_type}s.xml"
                note_part = None
                for rel in document.part.rels.values():
                    if hasattr(rel, 'target_part') and hasattr(rel.target_part, 'partname'):
                        if str(rel.target_part.partname) == part_name:
                            note_part = rel.target_part
                            break

                if note_part is None:
                    continue

                from lxml import etree
                root = etree.fromstring(note_part.blob)
                notes = root.findall(f'{{{w_ns}}}{note_type}')

                for note in notes:
                    # Skip separator/continuation notes
                    note_type_attr = note.get(f'{{{w_ns}}}type')
                    if note_type_attr in ('separator', 'continuationSeparator', 'continuationNotice'):
                        continue

                    note_id = note.get(f'{{{w_ns}}}id')
                    paragraphs = note.findall(f'{{{w_ns}}}p')

                    for para_idx, para_elem in enumerate(paragraphs):
                        text_parts = []
                        for t in para_elem.findall(f'.//{{{w_ns}}}t'):
                            if t.text:
                                text_parts.append(t.text)
                        full_text = ''.join(text_parts).strip()
                        if full_text:
                            segment = TextSegment(
                                id=str(segment_id),
                                text=full_text,
                                location=f"{note_type}_{note_id}_paragraph_{para_idx}",
                                metadata={
                                    "type": note_type,
                                    "note_id": note_id,
                                    "paragraph_idx": para_idx,
                                }
                            )
                            segments.append(segment)
                            segment_id += 1

            except Exception as e:
                self.logger.warning(f"Error extracting {note_type}s: {e}")

        return segments

    async def _extract_textbox_segments(
        self,
        document: Document,
        start_segment_id: int
    ) -> List[TextSegment]:
        """
        Extract text segments from text boxes and shapes.
        
        Text boxes in Word are stored as drawing elements in the document XML.
        """
        segments = []
        segment_id = start_segment_id
        
        # Access the document's XML to find text boxes
        try:
            body = document.element.body
            # Find all text box elements (w:txbxContent)
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
                'v': 'urn:schemas-microsoft-com:vml',
            }

            # Find text boxes using XPath (DrawingML + VML)
            textbox_contents = body.findall('.//wps:txbx/w:txbxContent', namespaces)
            seen_ids = {id(tc) for tc in textbox_contents}
            for tc in body.findall('.//v:textbox/w:txbxContent', namespaces):
                if id(tc) not in seen_ids:
                    textbox_contents.append(tc)
                    seen_ids.add(id(tc))
            
            w_ns = namespaces['w']
            for tb_idx, txbx_content in enumerate(textbox_contents):
                # Collect element IDs of paragraphs inside tables to avoid double-counting
                table_para_ids: Set[int] = set()
                tables = txbx_content.findall(f'{{{w_ns}}}tbl')
                for tbl_elem in tables:
                    for p in tbl_elem.findall(f'.//{{{w_ns}}}p'):
                        table_para_ids.add(id(p))

                # Extract direct paragraphs from text box (skip those inside tables)
                paragraphs = txbx_content.findall(f'.//{{{w_ns}}}p')
                direct_para_idx = 0
                for para_elem in paragraphs:
                    if id(para_elem) in table_para_ids:
                        continue
                    text_parts = []
                    text_runs = para_elem.findall(f'.//{{{w_ns}}}t')
                    for t in text_runs:
                        if t.text:
                            text_parts.append(t.text)

                    full_text = ''.join(text_parts).strip()
                    if full_text:
                        segment = TextSegment(
                            id=str(segment_id),
                            text=full_text,
                            location=f"textbox_{tb_idx}_paragraph_{direct_para_idx}",
                            metadata={
                                "type": "textbox",
                                "textbox_idx": tb_idx,
                                "paragraph_idx": direct_para_idx,
                            }
                        )
                        segments.append(segment)
                        segment_id += 1
                    direct_para_idx += 1

                # Extract tables from text box
                for tbl_idx, tbl_elem in enumerate(tables):
                    rows = tbl_elem.findall(f'{{{w_ns}}}tr')
                    for row_idx, row_elem in enumerate(rows):
                        cells = row_elem.findall(f'{{{w_ns}}}tc')
                        for col_idx, cell_elem in enumerate(cells):
                            cell_text_parts = []
                            for t in cell_elem.findall(f'.//{{{w_ns}}}t'):
                                if t.text:
                                    cell_text_parts.append(t.text)
                            cell_text = ''.join(cell_text_parts).strip()
                            if cell_text:
                                segment = TextSegment(
                                    id=str(segment_id),
                                    text=cell_text,
                                    location=f"textbox_{tb_idx}_table_{tbl_idx}_row_{row_idx}_col_{col_idx}",
                                    metadata={
                                        "type": "textbox_table_cell",
                                        "textbox_idx": tb_idx,
                                        "table_idx": tbl_idx,
                                        "row_idx": row_idx,
                                        "col_idx": col_idx,
                                    }
                                )
                                segments.append(segment)
                                segment_id += 1
        except Exception as e:
            self.logger.warning(f"Error extracting text boxes: {e}")
        
        return segments

    async def write_translated(
        self,
        file_path: Path,
        segments: List[TextSegment],
        translations: List[str],
        output_path: Path,
        output_mode: str = "replace"
    ) -> bool:
        """
        Write translated text back to Word document, preserving formatting.

        Args:
            file_path: Path to the original Word file
            segments: List of original text segments
            translations: List of translated texts (same order as segments)
            output_path: Path where the translated document should be saved
            output_mode: One of "replace", "append", "prepend", "interleave", "interleave_reverse" (default: "replace")
            
        Returns:
            True if writing succeeded, False otherwise
        """
        try:
            # Always load fresh from file to avoid shared state issues
            # during concurrent processing of multiple files
            document = await asyncio.to_thread(Document, str(file_path))
            
            # Create translation map with output mode applied
            translation_map = {}
            for seg, trans in zip(segments, translations):
                final_text = apply_output_mode(seg.text, trans, output_mode)
                translation_map[seg.id] = final_text
            
            # Process each segment
            for segment in segments:
                translation = translation_map.get(segment.id)
                if translation is None:
                    continue
                
                seg_type = segment.metadata.get("type")
                
                if seg_type == "paragraph":
                    await self._write_paragraph_translation(
                        document, segment, translation
                    )
                elif seg_type == "table_cell":
                    await self._write_table_cell_translation(
                        document, segment, translation
                    )
                elif seg_type == "textbox":
                    await self._write_textbox_translation(
                        document, segment, translation
                    )
                elif seg_type == "textbox_table_cell":
                    await self._write_textbox_table_cell_translation(
                        document, segment, translation
                    )
                elif seg_type == "sdt_paragraph":
                    self._write_sdt_translation(
                        document, segment, translation
                    )
                elif seg_type == "sdt_table_cell":
                    self._write_sdt_table_cell_translation(
                        document, segment, translation
                    )
                elif seg_type in ("footnote", "endnote"):
                    self._write_footnote_endnote_translation(
                        document, segment, translation
                    )
            
            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Save the document
            await asyncio.to_thread(document.save, str(output_path))
            
            self.logger.info(f"Saved translated Word document to: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to write translated Word document: {e}")
            return False

    async def _write_paragraph_translation(
        self,
        document: Document,
        segment: TextSegment,
        translation: str
    ) -> None:
        """Write translation to a paragraph while preserving formatting."""
        location_prefix = segment.metadata.get("location_prefix", "body")
        para_idx = segment.metadata.get("paragraph_idx", 0)
        header_footer = segment.metadata.get("header_footer")
        section_idx = segment.metadata.get("section_idx")
        
        # Find the target paragraph
        paragraph = None
        
        if header_footer == "header" and section_idx is not None:
            section = document.sections[section_idx]
            if para_idx < len(section.header.paragraphs):
                paragraph = section.header.paragraphs[para_idx]
        elif header_footer == "footer" and section_idx is not None:
            section = document.sections[section_idx]
            if para_idx < len(section.footer.paragraphs):
                paragraph = section.footer.paragraphs[para_idx]
        elif location_prefix == "body":
            if para_idx < len(document.paragraphs):
                paragraph = document.paragraphs[para_idx]
        
        if paragraph is None:
            self.logger.warning(f"Could not find paragraph for segment {segment.id}")
            return
        
        # Preserve formatting by updating runs
        runs_meta = segment.metadata.get("runs", [])
        
        if len(paragraph.runs) == 1 or not runs_meta:
            # Simple case: single run or no run metadata
            self._update_paragraph_text_simple(paragraph, translation)
        else:
            # Complex case: multiple runs with formatting
            self._update_paragraph_text_with_runs(paragraph, translation, runs_meta)
    
    def _clear_paragraph_content(self, paragraph: Paragraph) -> None:
        """Remove all content elements from paragraph XML, preserving paragraph properties (w:pPr)."""
        p_elem = paragraph._element
        pPr = p_elem.find(qn('w:pPr'))
        for child in list(p_elem):
            if child is not pPr:
                p_elem.remove(child)

    def _update_paragraph_text_simple(
        self,
        paragraph: Paragraph,
        translation: str
    ) -> None:
        """Update paragraph text when there's a single run or simple structure."""
        if paragraph.runs:
            # Clear all runs except the first, update first with full translation
            first_run = paragraph.runs[0]
            # Store formatting
            bold = first_run.bold
            italic = first_run.italic
            underline = first_run.underline
            font_name = first_run.font.name
            font_size = first_run.font.size
            
            # Clear paragraph (including hyperlinks/field codes for TOC entries)
            self._clear_paragraph_content(paragraph)

            # Add new run with translation
            new_run = paragraph.add_run(translation)
            new_run.bold = bold
            new_run.italic = italic
            new_run.underline = underline
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
        else:
            # Clear any non-run content (e.g. hyperlinks in TOC entries)
            self._clear_paragraph_content(paragraph)
            paragraph.add_run(translation)
    
    def _update_paragraph_text_with_runs(
        self,
        paragraph: Paragraph,
        translation: str,
        runs_meta: List[dict]
    ) -> None:
        """
        Update paragraph text while preserving run-level formatting.
        
        Strategy: Use the formatting from the first run for the entire translation,
        as we cannot reliably map character positions between languages.
        """
        if not runs_meta:
            self._update_paragraph_text_simple(paragraph, translation)
            return
        
        # Get formatting from first run
        first_run_meta = runs_meta[0]
        
        # Clear paragraph (including hyperlinks/field codes for TOC entries)
        self._clear_paragraph_content(paragraph)

        # Add new run with translation and preserved formatting
        new_run = paragraph.add_run(translation)
        
        if first_run_meta.get("bold"):
            new_run.bold = True
        if first_run_meta.get("italic"):
            new_run.italic = True
        if first_run_meta.get("underline"):
            new_run.underline = True
        if first_run_meta.get("font_name"):
            new_run.font.name = first_run_meta["font_name"]
        if first_run_meta.get("font_size"):
            new_run.font.size = Pt(first_run_meta["font_size"])
        if first_run_meta.get("font_color"):
            try:
                color_str = first_run_meta["font_color"]
                new_run.font.color.rgb = RGBColor.from_string(color_str)
            except Exception:
                pass

    async def _write_table_cell_translation(
        self,
        document: Document,
        segment: TextSegment,
        translation: str
    ) -> None:
        """Write translation to a table cell while preserving formatting."""
        table_idx = segment.metadata.get("table_idx", 0)
        row_idx = segment.metadata.get("row_idx", 0)
        col_idx = segment.metadata.get("col_idx", 0)
        location_prefix = segment.metadata.get("location_prefix", "")
        header_footer = segment.metadata.get("header_footer")
        section_idx = segment.metadata.get("section_idx")
        
        # Find the correct table based on location
        table = self._find_table(
            document, table_idx, location_prefix, header_footer, section_idx
        )
        
        if table is None:
            self.logger.warning(f"Could not find table for segment {segment.id}")
            return
        
        if row_idx >= len(table.rows):
            self.logger.warning(f"Row index {row_idx} out of range in table {table_idx}")
            return
        
        row = table.rows[row_idx]
        
        if col_idx >= len(row.cells):
            self.logger.warning(f"Column index {col_idx} out of range in table {table_idx}")
            return
        
        # Use index-based lookup - this works correctly for both merged and non-merged cells
        # For merged cells, row.cells[col_idx] returns the merged cell which is what we want
        cell = row.cells[col_idx]
        
        # Get paragraph metadata if available
        paragraphs_meta = segment.metadata.get("paragraphs", [])
        
        if cell.paragraphs:
            # Split translation into lines to preserve multi-paragraph structure
            translated_lines = translation.split('\n') if '\n' in translation else [translation]
            cell_paras = cell.paragraphs

            for i, para in enumerate(cell_paras):
                if i < len(translated_lines):
                    # Find matching formatting from paragraphs_meta
                    runs_meta = []
                    if paragraphs_meta:
                        # Match by para_idx stored in metadata
                        for pm in paragraphs_meta:
                            if pm.get("para_idx") == i:
                                runs_meta = pm.get("runs", [])
                                break
                        # Fallback: use positional match
                        if not runs_meta and i < len(paragraphs_meta):
                            runs_meta = paragraphs_meta[i].get("runs", [])

                    if runs_meta:
                        self._update_paragraph_text_with_runs(para, translated_lines[i], runs_meta)
                    else:
                        self._update_paragraph_text_simple(para, translated_lines[i])
                else:
                    # More paragraphs than translated lines -- clear extras
                    para.clear()

            # If more translated lines than paragraphs, append overflow to last paragraph
            if len(translated_lines) > len(cell_paras):
                overflow = '\n'.join(translated_lines[len(cell_paras):])
                last_para = cell_paras[-1]
                current_text = last_para.text
                combined = current_text + '\n' + overflow if current_text else overflow
                self._update_paragraph_text_simple(last_para, combined)
    
    def _find_table(
        self,
        document: Document,
        table_idx: int,
        location_prefix: str,
        header_footer: Optional[str],
        section_idx: Optional[int]
    ) -> Optional[Table]:
        """
        Find the correct table based on location metadata.
        
        Handles tables in:
        - Document body
        - Headers/footers
        - Nested within other table cells
        """
        # Handle tables in headers/footers
        if header_footer and section_idx is not None:
            if section_idx >= len(document.sections):
                return None
            section = document.sections[section_idx]
            
            if header_footer == "header":
                container = section.header
            else:
                container = section.footer
            
            if table_idx >= len(container.tables):
                return None
            return container.tables[table_idx]
        
        # Handle nested tables (location_prefix contains "nested_")
        if "nested_" in location_prefix:
            return self._find_nested_table(document, location_prefix, table_idx)
        
        # Regular body table
        if table_idx >= len(document.tables):
            return None
        return document.tables[table_idx]
    
    def _find_nested_table(
        self,
        document: Document,
        location_prefix: str,
        nested_table_idx: int
    ) -> Optional[Table]:
        """
        Find a nested table by parsing the location prefix.
        
        Location prefix format: "table_X_row_Y_col_Z_nested_"
        """
        try:
            # Parse the location prefix to find parent cell
            # Example: "table_0_row_1_col_2_nested_"
            parts = location_prefix.rstrip("_").split("_")
            
            # Find parent table index
            parent_table_idx = None
            parent_row_idx = None
            parent_col_idx = None
            
            i = 0
            while i < len(parts):
                if parts[i] == "table" and i + 1 < len(parts):
                    parent_table_idx = int(parts[i + 1])
                    i += 2
                elif parts[i] == "row" and i + 1 < len(parts):
                    parent_row_idx = int(parts[i + 1])
                    i += 2
                elif parts[i] == "col" and i + 1 < len(parts):
                    parent_col_idx = int(parts[i + 1])
                    i += 2
                elif parts[i] == "nested":
                    break
                else:
                    i += 1
            
            if parent_table_idx is None or parent_row_idx is None or parent_col_idx is None:
                return None
            
            # Navigate to parent table
            if parent_table_idx >= len(document.tables):
                return None
            parent_table = document.tables[parent_table_idx]
            
            if parent_row_idx >= len(parent_table.rows):
                return None
            parent_row = parent_table.rows[parent_row_idx]
            
            if parent_col_idx >= len(parent_row.cells):
                return None
            parent_cell = parent_row.cells[parent_col_idx]
            
            # Get nested table
            if nested_table_idx >= len(parent_cell.tables):
                return None
            return parent_cell.tables[nested_table_idx]
            
        except (ValueError, IndexError) as e:
            self.logger.warning(f"Error parsing nested table location: {e}")
            return None
    
    async def _write_textbox_translation(
        self,
        document: Document,
        segment: TextSegment,
        translation: str
    ) -> None:
        """Write translation to a text box."""
        textbox_idx = segment.metadata.get("textbox_idx", 0)
        para_idx = segment.metadata.get("paragraph_idx", 0)
        
        try:
            body = document.element.body
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
                'v': 'urn:schemas-microsoft-com:vml',
            }

            textbox_contents = body.findall('.//wps:txbx/w:txbxContent', namespaces)
            seen_ids = {id(tc) for tc in textbox_contents}
            for tc in body.findall('.//v:textbox/w:txbxContent', namespaces):
                if id(tc) not in seen_ids:
                    textbox_contents.append(tc)
                    seen_ids.add(id(tc))
            
            if textbox_idx >= len(textbox_contents):
                self.logger.warning(f"Textbox index {textbox_idx} out of range")
                return
            
            txbx_content = textbox_contents[textbox_idx]
            w_ns = namespaces['w']

            # Collect paragraphs inside tables to skip them (match extraction logic)
            table_para_ids: Set[int] = set()
            for tbl in txbx_content.findall(f'{{{w_ns}}}tbl'):
                for p in tbl.findall(f'.//{{{w_ns}}}p'):
                    table_para_ids.add(id(p))

            # Get only direct (non-table) paragraphs
            direct_paragraphs = [
                p for p in txbx_content.findall(f'.//{{{w_ns}}}p')
                if id(p) not in table_para_ids
            ]

            if para_idx >= len(direct_paragraphs):
                self.logger.warning(f"Paragraph index {para_idx} out of range in textbox {textbox_idx}")
                return

            para_elem = direct_paragraphs[para_idx]
            
            # Find all text runs and update them
            text_runs = para_elem.findall('.//w:t', namespaces)
            
            if text_runs:
                # Put all translation in first text run, clear others
                text_runs[0].text = translation
                for t in text_runs[1:]:
                    t.text = ""
                    
        except Exception as e:
            self.logger.warning(f"Error writing to textbox: {e}")

    async def _write_textbox_table_cell_translation(
        self,
        document: Document,
        segment: TextSegment,
        translation: str
    ) -> None:
        """Write translation to a table cell inside a text box."""
        textbox_idx = segment.metadata.get("textbox_idx", 0)
        table_idx = segment.metadata.get("table_idx", 0)
        row_idx = segment.metadata.get("row_idx", 0)
        col_idx = segment.metadata.get("col_idx", 0)

        try:
            body = document.element.body
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
                'v': 'urn:schemas-microsoft-com:vml',
            }
            w_ns = namespaces['w']

            textbox_contents = body.findall('.//wps:txbx/w:txbxContent', namespaces)
            seen_ids = {id(tc) for tc in textbox_contents}
            for tc in body.findall('.//v:textbox/w:txbxContent', namespaces):
                if id(tc) not in seen_ids:
                    textbox_contents.append(tc)
                    seen_ids.add(id(tc))

            if textbox_idx >= len(textbox_contents):
                self.logger.warning(f"Textbox index {textbox_idx} out of range")
                return

            txbx_content = textbox_contents[textbox_idx]
            tables = txbx_content.findall(f'{{{w_ns}}}tbl')

            if table_idx >= len(tables):
                self.logger.warning(f"Table index {table_idx} out of range in textbox {textbox_idx}")
                return

            tbl_elem = tables[table_idx]
            rows = tbl_elem.findall(f'{{{w_ns}}}tr')
            if row_idx >= len(rows):
                return
            cells = rows[row_idx].findall(f'{{{w_ns}}}tc')
            if col_idx >= len(cells):
                return

            cell_elem = cells[col_idx]
            text_runs = cell_elem.findall(f'.//{{{w_ns}}}t')
            if text_runs:
                text_runs[0].text = translation
                for t in text_runs[1:]:
                    t.text = ""

        except Exception as e:
            self.logger.warning(f"Error writing to textbox table cell: {e}")

    def _write_sdt_translation(
        self,
        document: Document,
        segment: TextSegment,
        translation: str
    ) -> None:
        """Write translation to a paragraph inside an SDT content control."""
        sdt_idx = segment.metadata.get("sdt_idx", 0)
        para_idx = segment.metadata.get("paragraph_idx", 0)
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        try:
            body = document.element.body
            sdt_elements = [
                child for child in body
                if child.tag == f'{{{w_ns}}}sdt'
            ]

            if sdt_idx >= len(sdt_elements):
                self.logger.warning(f"SDT index {sdt_idx} out of range")
                return

            sdt_content = sdt_elements[sdt_idx].find(f'{{{w_ns}}}sdtContent')
            if sdt_content is None:
                return

            paragraphs = [
                elem for elem in sdt_content
                if elem.tag == f'{{{w_ns}}}p'
            ]

            if para_idx >= len(paragraphs):
                self.logger.warning(f"Paragraph index {para_idx} out of range in SDT {sdt_idx}")
                return

            para_elem = paragraphs[para_idx]
            text_runs = para_elem.findall(f'.//{{{w_ns}}}t')
            if text_runs:
                text_runs[0].text = translation
                for t in text_runs[1:]:
                    t.text = ""
            else:
                # No existing runs; add one
                from lxml import etree
                r_elem = etree.SubElement(para_elem, f'{{{w_ns}}}r')
                t_elem = etree.SubElement(r_elem, f'{{{w_ns}}}t')
                t_elem.text = translation

        except Exception as e:
            self.logger.warning(f"Error writing to SDT paragraph: {e}")

    def _write_sdt_table_cell_translation(
        self,
        document: Document,
        segment: TextSegment,
        translation: str
    ) -> None:
        """Write translation to a table cell inside an SDT content control."""
        sdt_idx = segment.metadata.get("sdt_idx", 0)
        table_idx = segment.metadata.get("table_idx", 0)
        row_idx = segment.metadata.get("row_idx", 0)
        col_idx = segment.metadata.get("col_idx", 0)
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        try:
            body = document.element.body
            sdt_elements = [
                child for child in body
                if child.tag == f'{{{w_ns}}}sdt'
            ]

            if sdt_idx >= len(sdt_elements):
                return

            sdt_content = sdt_elements[sdt_idx].find(f'{{{w_ns}}}sdtContent')
            if sdt_content is None:
                return

            tables = [
                elem for elem in sdt_content
                if elem.tag == f'{{{w_ns}}}tbl'
            ]

            if table_idx >= len(tables):
                return

            tbl_elem = tables[table_idx]
            rows = tbl_elem.findall(f'{{{w_ns}}}tr')
            if row_idx >= len(rows):
                return
            cells = rows[row_idx].findall(f'{{{w_ns}}}tc')
            if col_idx >= len(cells):
                return

            cell_elem = cells[col_idx]
            text_runs = cell_elem.findall(f'.//{{{w_ns}}}t')
            if text_runs:
                text_runs[0].text = translation
                for t in text_runs[1:]:
                    t.text = ""

        except Exception as e:
            self.logger.warning(f"Error writing to SDT table cell: {e}")

    def _write_footnote_endnote_translation(
        self,
        document: Document,
        segment: TextSegment,
        translation: str
    ) -> None:
        """Write translation to a footnote or endnote."""
        note_type = segment.metadata.get("type")  # "footnote" or "endnote"
        note_id = segment.metadata.get("note_id")
        para_idx = segment.metadata.get("paragraph_idx", 0)
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        try:
            part_name = f"/word/{note_type}s.xml"
            note_part = None
            for rel in document.part.rels.values():
                if hasattr(rel, 'target_part') and hasattr(rel.target_part, 'partname'):
                    if str(rel.target_part.partname) == part_name:
                        note_part = rel.target_part
                        break

            if note_part is None:
                return

            from lxml import etree
            root = etree.fromstring(note_part.blob)
            notes = root.findall(f'{{{w_ns}}}{note_type}')

            for note in notes:
                if note.get(f'{{{w_ns}}}id') == note_id:
                    paragraphs = note.findall(f'{{{w_ns}}}p')
                    if para_idx < len(paragraphs):
                        para_elem = paragraphs[para_idx]
                        text_runs = para_elem.findall(f'.//{{{w_ns}}}t')
                        if text_runs:
                            text_runs[0].text = translation
                            for t in text_runs[1:]:
                                t.text = ""
                    break

            # Write modified XML back to part
            note_part._blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

        except Exception as e:
            self.logger.warning(f"Error writing to {note_type}: {e}")

    async def validate_file(self, file_path: Path) -> tuple[bool, Optional[str]]:
        """
        Validate that the Word file can be processed.
        
        Checks for:
        - File existence
        - File readability (not password protected or corrupted)
        - Valid .docx format
        
        Args:
            file_path: Path to the Word file to validate
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not file_path.exists():
            return False, f"File not found: {file_path}"
        
        if file_path.suffix.lower() not in self.supported_extensions:
            return False, f"Unsupported file format: {file_path.suffix}"
        
        try:
            # Try to load the document to check for corruption/password protection
            document = await asyncio.to_thread(Document, str(file_path))
            
            # Check if document has any content
            has_content = (
                len(document.paragraphs) > 0 or
                len(document.tables) > 0
            )
            
            if not has_content:
                # Check headers/footers for content
                for section in document.sections:
                    if section.header.paragraphs or section.footer.paragraphs:
                        has_content = True
                        break
            
            if not has_content:
                return False, "Word document contains no content."
            
            return True, None
            
        except PackageNotFoundError:
            return False, "This file appears to be corrupted and cannot be read."
        except Exception as e:
            error_msg = str(e).lower()
            if "password" in error_msg or "encrypted" in error_msg:
                return False, "This file is password protected. Please remove the password and try again."
            if "package" in error_msg or "zip" in error_msg:
                return False, "This file appears to be corrupted and cannot be read."
            return False, f"Failed to read Word file: {str(e)}"
